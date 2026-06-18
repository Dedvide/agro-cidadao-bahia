"""
Extrator com IA — o coração técnico do SIPA-Bahia AI.

Recebe o TEXTO de um laudo heterogêneo (Excel/Word/PDF/CSV/máquina já lidos
como texto) e devolve um objeto `Laudo` estruturado e validado.

Estratégia: Claude com SAÍDA ESTRUTURADA (tool use forçado) contra o esquema
canônico. O modelo é obrigado a preencher o JSON Schema — a validação acontece
na camada da tool, então o modelo "tenta de novo" sozinho se errar o formato.
Prompt caching nas instruções (elas repetem a cada laudo → barateia muito).

Soberania de dados: troque o provedor por um modelo open-source local quando
for para produção. A interface (texto -> Laudo) permanece a mesma.
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

from schema import LAUDO_JSON_SCHEMA, Laudo

# Modelo configurável por variável de ambiente.
MODELO = os.environ.get("CETAB_MODELO", "claude-sonnet-4-6")

INSTRUCOES = """Você é o extrator de laudos laboratoriais do CETAB (Centro \
de Tecnologias Agropecuárias da Bahia). Sua tarefa é ler o laudo (que pode vir \
em qualquer layout: texto corrido, tabela, CSV, planilha exportada) e preencher \
o esquema canônico chamando a ferramenta `registrar_laudo`.

REGRAS GERAIS:
- Converta vírgula decimal para ponto (5,2 -> 5.2).
- Datas sempre em AAAA-MM-DD. Se o laudo traz 12/03/2026, vire 2026-03-12.
- NÃO INVENTE dados. Se um campo não existe ou está ilegível, deixe null e
  registre o nome dele em `campos_incertos`.
- `confianca_geral`: 1.0 quando tudo está claro; abaixo de 0.7 quando o layout
  for ambíguo, faltar unidade, ou você tiver dúvida sobre algum valor.
- Capture TODOS os itens presentes, um por `resultados`.
- Infira `solicitante_tipo`, `tipo_amostra` e `categoria_analise` pelo contexto.

CADA RESULTADO TEM UM `tipo` — escolha conforme a natureza do laudo:
- "quantitativo" (química: solo, água, foliar, resíduos, metais, alimentos):
  preencha analito (minúsculo, sem acento, underscore: "Fósforo (P)"->"fosforo",
  "pH"->"ph", "Matéria orgânica"->"materia_organica"), valor, unidade, metodo.
- "diagnostico" (biologia molecular / sanidade: HLB, CVC, viroses):
  preencha alvo (patógeno, minúsculo: "hlb","cvc","cabmv"),
  resultado_diagnostico ("positivo"/"negativo"/"inconclusivo") e ct se houver.
- "identificacao" (entomologia / pragas / fungos):
  preencha taxon (espécie), contagem (nº) e estagio se houver.

`categoria_analise` (minúsculo): fertilidade_solo, analise_agua, analise_foliar,
residuos_agrotoxicos, metais_pesados, qualidade_alimentos, diagnostico_molecular,
monitoramento_pragas, fitopatologia."""

TOOL = {
    "name": "registrar_laudo",
    "description": "Registra os dados extraídos do laudo no esquema canônico do CETAB.",
    "input_schema": LAUDO_JSON_SCHEMA,
}


def extrair(texto_documento: str) -> Laudo:
    """Extrai um Laudo estruturado a partir do texto bruto de um documento."""
    import anthropic  # import tardio para a mensagem de erro de dependência ser clara

    client = anthropic.Anthropic()  # usa ANTHROPIC_API_KEY do ambiente
    resp = client.messages.create(
        model=MODELO,
        max_tokens=2048,
        system=[
            {
                "type": "text",
                "text": INSTRUCOES,
                "cache_control": {"type": "ephemeral"},  # cacheia as instruções
            }
        ],
        tools=[TOOL],
        tool_choice={"type": "tool", "name": "registrar_laudo"},
        messages=[
            {
                "role": "user",
                "content": f"Extraia os dados deste laudo:\n\n<laudo>\n{texto_documento}\n</laudo>",
            }
        ],
    )
    for bloco in resp.content:
        if bloco.type == "tool_use":
            return Laudo.model_validate(bloco.input)
    raise RuntimeError("O modelo não retornou uma chamada de ferramenta (tool_use).")


def carregar_texto(caminho: Path) -> str:
    """Lê o documento como texto, despachando por extensão.

    .txt/.csv → direto · .docx → python-docx · .pdf → pypdf.
    Em produção, .pdf escaneado/imagem usa docling/unstructured (+ OCR).
    """
    sufixo = caminho.suffix.lower()
    if sufixo in (".txt", ".csv"):
        return caminho.read_text(encoding="utf-8")
    if sufixo == ".docx":
        return _ler_docx(caminho)
    if sufixo == ".pdf":
        return _ler_pdf(caminho)
    raise ValueError(f"Formato não suportado no protótipo: {sufixo}")


def _ler_docx(caminho: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError("Instale: pip install python-docx") from e
    doc = docx.Document(str(caminho))
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabela in doc.tables:  # tabelas viram linhas " | "-separadas
        for linha in tabela.rows:
            partes.append(" | ".join(c.text for c in linha.cells))
    return "\n".join(partes)


def _ler_pdf(caminho: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("Instale: pip install pypdf") from e
    reader = PdfReader(str(caminho))
    return "\n".join((pagina.extract_text() or "") for pagina in reader.pages)


def _main(argv: list[str]) -> int:
    if not argv:
        print("Uso: python extrator.py <caminho_do_laudo>")
        return 2
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print(
            "ERRO: defina a variável de ambiente ANTHROPIC_API_KEY.\n"
            "PowerShell:  $env:ANTHROPIC_API_KEY = 'sua-chave'"
        )
        return 1

    caminho = Path(argv[0])
    texto = carregar_texto(caminho)
    laudo = extrair(texto)
    print(laudo.model_dump_json(indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(_main(sys.argv[1:]))
