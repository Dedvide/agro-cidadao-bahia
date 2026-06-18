"""Gera um laudo .docx (com tabela) e testa o loader carregar_texto sobre ele."""
from pathlib import Path

import docx

from extrator import carregar_texto

PASTA = Path(__file__).parent
ALVO = PASTA / "exemplos" / "laudo_solo_03.docx"

# --- gera um .docx realista com parágrafos + tabela ---
doc = docx.Document()
doc.add_heading("CETAB - Laudo de Análise de Solo", level=1)
doc.add_paragraph("Laudo nº: 2026/0488")
doc.add_paragraph("Solicitante: Maria das Graças (Cooperativa)")
doc.add_paragraph("Município: Muritiba - BA")
doc.add_paragraph("Data da coleta: 18/03/2026")

tabela = doc.add_table(rows=1, cols=3)
tabela.rows[0].cells[0].text = "Parâmetro"
tabela.rows[0].cells[1].text = "Resultado"
tabela.rows[0].cells[2].text = "Unidade"
for param, res, uni in [
    ("pH", "5,8", "-"),
    ("Fósforo", "14", "mg/dm3"),
    ("Potássio", "60", "mg/dm3"),
]:
    linha = tabela.add_row().cells
    linha[0].text, linha[1].text, linha[2].text = param, res, uni

doc.save(str(ALVO))
print(f"Gerado: {ALVO.name}")

# --- testa o loader ---
texto = carregar_texto(ALVO)
print("=" * 50)
print("TEXTO EXTRAÍDO DO .DOCX (vai para a IA):")
print("=" * 50)
print(texto)
print("=" * 50)
assert "Fósforo" in texto and "mg/dm3" in texto, "Falha: tabela não extraída"
print("Loader .docx (parágrafos + tabela): OK")
